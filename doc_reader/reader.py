from docx import Document
from PyQt6.QAxContainer import QAxObject


a = Document('C:/Users/detec/OneDrive/Документы/выпускной.docx')
for i in Document('C:/Users/detec/OneDrive/Документы/выпускной.docx').paragraphs:
    print(i.text)
a = QAxObject()
